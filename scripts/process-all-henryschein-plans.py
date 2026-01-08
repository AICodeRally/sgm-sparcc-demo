#!/usr/bin/env python3
"""
Process All Henry Schein Compensation Plans
Converts PDF and DOCX files to structured markdown format
"""

import os
import json
import re
from pathlib import Path
from typing import Dict, List, Optional

# Try importing PDF libraries
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False
    print("Warning: PyPDF2 not installed. Trying pdfplumber...")

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False
    print("Warning: pdfplumber not installed")

# Try importing DOCX library
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Warning: python-docx not installed")

# Configuration
BASE_PATH = Path("/Users/toddlebaron/Documents/SPM/clients/HenrySchein/HS_Comp_Plans")
RAW_PATH = BASE_PATH / "raw"
PROCESSED_PATH = BASE_PATH / "processed"
MANIFEST_PATH = BASE_PATH / "plans-manifest.json"

def load_manifest() -> Dict:
    """Load the plans manifest"""
    with open(MANIFEST_PATH, 'r') as f:
        return json.load(f)

def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extract text from PDF using available library"""
    text = ""

    # Try pdfplumber first (better quality)
    if HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            return text
        except Exception as e:
            print(f"  pdfplumber failed: {e}, trying PyPDF2...")

    # Fall back to PyPDF2
    if HAS_PYPDF2:
        try:
            with open(pdf_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            return text
        except Exception as e:
            print(f"  PyPDF2 failed: {e}")
            return ""

    print(f"  No PDF library available to extract text")
    return ""

def extract_text_from_docx(docx_path: Path) -> str:
    """Extract text from DOCX file"""
    if not HAS_DOCX:
        print(f"  python-docx not available")
        return ""

    try:
        doc = Document(docx_path)
        text = ""

        for para in doc.paragraphs:
            # Detect headings and add markdown formatting
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.replace('Heading ', ''))
                prefix = '#' * min(level + 1, 6)  # Cap at H6
                text += f"\n{prefix} {para.text}\n\n"
            else:
                if para.text.strip():
                    text += para.text + "\n\n"

        # Extract tables
        for table in doc.tables:
            text += "\n"
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                text += "| " + " | ".join(cells) + " |\n"
            text += "\n"

        return text
    except Exception as e:
        print(f"  Error extracting DOCX: {e}")
        return ""

def clean_text(text: str) -> str:
    """Clean and normalize extracted text"""
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Remove page numbers (common patterns)
    text = re.sub(r'\n\s*Page \d+\s*\n', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'\n\s*\d+\s*\n', '\n', text)

    # Fix common OCR issues
    text = text.replace('ﬁ', 'fi')
    text = text.replace('ﬂ', 'fl')
    text = text.replace('–', '-')
    text = text.replace('—', '-')

    # Normalize quotes
    text = text.replace('"', '"').replace('"', '"')
    text = text.replace(''', "'").replace(''', "'")

    return text.strip()

def detect_sections(text: str) -> str:
    """Detect and markup sections as markdown headers"""
    lines = text.split('\n')
    result = []

    for line in lines:
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            result.append(line)
            continue

        # Detect section headers (common patterns)
        # Pattern 1: "1. Section Title" or "1.0 Section Title"
        if re.match(r'^\d+\.[\d\s]*[A-Z][^.]*$', stripped):
            result.append(f"\n## {stripped}\n")
            continue

        # Pattern 2: "a) Subsection" or "a. Subsection"
        if re.match(r'^[a-z][\)\.]\s+[A-Z]', stripped):
            result.append(f"\n### {stripped}\n")
            continue

        # Pattern 3: ALL CAPS titles (but not too long)
        if stripped.isupper() and len(stripped) < 80 and not re.search(r'\d{4}', stripped):
            result.append(f"\n## {stripped.title()}\n")
            continue

        # Keep line as-is
        result.append(line)

    return '\n'.join(result)

def create_markdown_plan(plan_info: Dict, content: str) -> str:
    """Create structured markdown from plan info and content"""
    md = f"""# {plan_info['title']}

**Plan Code:** {plan_info['documentCode']}
**Division:** {plan_info['division']}
**Role:** {plan_info['role']}
**Effective Date:** {plan_info['effectiveDate']}
**Plan Year:** {plan_info['planYear']}
**Version:** {plan_info['version']}
**Status:** {plan_info['status']}
**Company:** Henry Schein

---

"""

    # Add cleaned and formatted content
    cleaned_content = clean_text(content)
    formatted_content = detect_sections(cleaned_content)

    md += formatted_content

    return md

def process_plan(plan_info: Dict) -> bool:
    """Process a single plan"""
    plan_code = plan_info['documentCode']
    filename = plan_info['filename']

    print(f"\nProcessing: {plan_code}")
    print(f"  File: {filename}")

    # Check if already processed
    output_path = PROCESSED_PATH / f"{plan_code}.md"
    if output_path.exists():
        print(f"  ✓ Already exists, skipping")
        return True

    # Find source file
    source_path = RAW_PATH / filename
    if not source_path.exists():
        print(f"  ✗ Source file not found: {source_path}")
        return False

    # Extract text based on file type
    file_ext = source_path.suffix.lower()

    if file_ext == '.pdf':
        if not (HAS_PDFPLUMBER or HAS_PYPDF2):
            print(f"  ✗ No PDF library available")
            return False
        content = extract_text_from_pdf(source_path)
    elif file_ext in ['.docx', '.doc']:
        if not HAS_DOCX:
            print(f"  ✗ python-docx not available")
            return False
        content = extract_text_from_docx(source_path)
    else:
        print(f"  ✗ Unsupported file type: {file_ext}")
        return False

    if not content.strip():
        print(f"  ✗ No content extracted")
        return False

    # Create markdown
    markdown = create_markdown_plan(plan_info, content)

    # Save to processed folder
    PROCESSED_PATH.mkdir(parents=True, exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(markdown)

    print(f"  ✓ Saved to: {output_path}")
    print(f"  ✓ Content length: {len(markdown)} characters")

    return True

def main():
    """Main processing function"""
    print("=" * 80)
    print("Henry Schein Compensation Plans - Batch Processor")
    print("=" * 80)

    # Check dependencies
    print("\nDependency Check:")
    print(f"  PyPDF2: {'✓' if HAS_PYPDF2 else '✗'}")
    print(f"  pdfplumber: {'✓' if HAS_PDFPLUMBER else '✗'}")
    print(f"  python-docx: {'✓' if HAS_DOCX else '✗'}")

    if not (HAS_PYPDF2 or HAS_PDFPLUMBER):
        print("\n⚠️  No PDF library found. Install with:")
        print("     pip install pdfplumber PyPDF2")

    if not HAS_DOCX:
        print("\n⚠️  python-docx not found. Install with:")
        print("     pip install python-docx")

    # Load manifest
    print(f"\nLoading manifest from: {MANIFEST_PATH}")
    manifest = load_manifest()
    plans = manifest['plans']

    print(f"Found {len(plans)} plans to process")

    # Process each plan
    success_count = 0
    skip_count = 0
    fail_count = 0

    for plan in plans:
        result = process_plan(plan)
        if result:
            if (PROCESSED_PATH / f"{plan['documentCode']}.md").exists():
                success_count += 1
        else:
            fail_count += 1

    # Summary
    print("\n" + "=" * 80)
    print("PROCESSING COMPLETE")
    print("=" * 80)
    print(f"✓ Successfully processed: {success_count}")
    print(f"✗ Failed: {fail_count}")
    print(f"Total plans: {len(plans)}")
    print(f"\nOutput directory: {PROCESSED_PATH}")
    print("=" * 80)

if __name__ == "__main__":
    main()
